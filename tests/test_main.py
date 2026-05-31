"""
Tests for the FastAPI main application.
"""

import io
import os
import tempfile
from pathlib import Path

import pytest
import fitz  # PyMuPDF
from docx import Document
from fastapi.testclient import TestClient

# Add backend directory to Python path
import sys
sys.path.insert(0, str(Path(__file__).parent.parent / "backend"))

from main import app, documents, UPLOAD_DIR


@pytest.fixture
def client():
    """
    Create a FastAPI test client.

    Yields:
        TestClient: FastAPI test client instance.
    """
    # Clear the in-memory documents store before each test
    documents.clear()
    return TestClient(app)


@pytest.fixture
def sample_pdf_bytes():
    """
    Create sample PDF file content as bytes.

    Returns:
        tuple: (filename, bytes content)
    """
    # Create a PDF in memory
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "This is a test PDF document.\nIt has multiple lines of text for testing.")

    # Save to bytes
    pdf_bytes = doc.tobytes()
    doc.close()

    return ("test_document.pdf", pdf_bytes)


@pytest.fixture
def sample_docx_bytes():
    """
    Create sample DOCX file content as bytes.

    Returns:
        tuple: (filename, bytes content)
    """
    # Create a DOCX in memory
    doc = Document()
    doc.add_paragraph("This is a test DOCX document.")
    doc.add_paragraph("It has multiple paragraphs for testing.")

    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    return ("test_document.docx", docx_bytes)


class TestHealthEndpoint:
    """Tests for the /health endpoint."""

    def test_health_check(self, client):
        """Test that health endpoint returns ok status."""
        response = client.get("/health")

        assert response.status_code == 200
        assert response.json() == {"status": "ok"}


class TestUploadEndpoint:
    """Tests for the /upload endpoint."""

    def test_upload_valid_pdf(self, client, sample_pdf_bytes):
        """Test uploading a valid PDF file."""
        filename, content = sample_pdf_bytes

        response = client.post(
            "/upload",
            files={"file": (filename, content, "application/pdf")}
        )

        assert response.status_code == 200
        data = response.json()

        # Check response structure
        assert "filename" in data
        assert "document_id" in data
        assert "char_count" in data
        assert "preview" in data

        # Check response values
        assert data["filename"] == filename
        assert len(data["document_id"]) == 36  # UUID4 format
        assert data["char_count"] > 0
        assert "test PDF document" in data["preview"]
        assert len(data["preview"]) <= 300  # Preview is max 300 chars

        # Verify document was stored in memory
        doc_id = data["document_id"]
        assert doc_id in documents
        assert documents[doc_id]["filename"] == filename
        assert documents[doc_id]["char_count"] == data["char_count"]

    def test_upload_valid_docx(self, client, sample_docx_bytes):
        """Test uploading a valid DOCX file."""
        filename, content = sample_docx_bytes

        response = client.post(
            "/upload",
            files={"file": (filename, content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )

        assert response.status_code == 200
        data = response.json()

        # Check response structure
        assert "filename" in data
        assert "document_id" in data
        assert "char_count" in data
        assert "preview" in data

        # Check response values
        assert data["filename"] == filename
        assert len(data["document_id"]) == 36  # UUID4 format
        assert data["char_count"] > 0
        assert "test DOCX document" in data["preview"]

    def test_upload_invalid_file_type(self, client):
        """Test that uploading invalid file type returns 400 error."""
        # Try to upload a .txt file
        response = client.post(
            "/upload",
            files={"file": ("document.txt", b"Some text content", "text/plain")}
        )

        assert response.status_code == 400
        data = response.json()
        assert "detail" in data
        assert ".txt" in data["detail"]
        assert "not allowed" in data["detail"].lower()

    def test_upload_invalid_extension_image(self, client):
        """Test that uploading image file returns 400 error."""
        response = client.post(
            "/upload",
            files={"file": ("image.jpg", b"fake image data", "image/jpeg")}
        )

        assert response.status_code == 400
        data = response.json()
        assert "detail" in data
        assert ".jpg" in data["detail"]

    def test_upload_creates_unique_filenames(self, client, sample_pdf_bytes):
        """Test that multiple uploads with same filename don't collide."""
        filename, content = sample_pdf_bytes

        # Upload same file twice
        response1 = client.post(
            "/upload",
            files={"file": (filename, content, "application/pdf")}
        )
        response2 = client.post(
            "/upload",
            files={"file": (filename, content, "application/pdf")}
        )

        assert response1.status_code == 200
        assert response2.status_code == 200

        # Document IDs should be different
        doc_id1 = response1.json()["document_id"]
        doc_id2 = response2.json()["document_id"]
        assert doc_id1 != doc_id2

        # Both files should exist with different paths
        assert documents[doc_id1]["filepath"] != documents[doc_id2]["filepath"]

    def test_upload_empty_pdf(self, client):
        """Test uploading an empty PDF file."""
        # Create an empty PDF
        doc = fitz.open()
        doc.new_page()  # Empty page
        pdf_bytes = doc.tobytes()
        doc.close()

        response = client.post(
            "/upload",
            files={"file": ("empty.pdf", pdf_bytes, "application/pdf")}
        )

        assert response.status_code == 200
        data = response.json()

        # Empty PDF should have 0 character count
        assert data["char_count"] == 0
        assert data["preview"] == ""

    def test_upload_preview_length(self, client):
        """Test that preview is truncated to 300 characters."""
        # Create a PDF with long text
        doc = fitz.open()
        page = doc.new_page()
        long_text = "A" * 500  # 500 characters
        page.insert_text((72, 72), long_text)
        pdf_bytes = doc.tobytes()
        doc.close()

        response = client.post(
            "/upload",
            files={"file": ("long.pdf", pdf_bytes, "application/pdf")}
        )

        assert response.status_code == 200
        data = response.json()

        # Preview should be exactly 300 chars
        assert len(data["preview"]) == 300
        # But full text should be longer
        assert data["char_count"] > 300

    def test_upload_file_saved_to_disk(self, client, sample_pdf_bytes):
        """Test that uploaded file is actually saved to disk."""
        filename, content = sample_pdf_bytes

        response = client.post(
            "/upload",
            files={"file": (filename, content, "application/pdf")}
        )

        assert response.status_code == 200
        doc_id = response.json()["document_id"]

        # Check file exists on disk
        filepath = Path(documents[doc_id]["filepath"])
        assert filepath.exists()
        assert filepath.parent == UPLOAD_DIR

        # Cleanup
        if filepath.exists():
            filepath.unlink()


class TestStartupBehavior:
    """Tests for application startup behavior."""

    def test_uploads_directory_created(self, client):
        """Test that uploads directory is created on startup."""
        # The directory should exist after client initialization
        assert UPLOAD_DIR.exists()
        assert UPLOAD_DIR.is_dir()
