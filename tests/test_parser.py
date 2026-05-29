"""
Tests for the document parser module.
"""

import os
import tempfile
from pathlib import Path

import pytest
import fitz  # PyMuPDF
from docx import Document

# Add backend directory to Python path
import sys
sys.path.insert(0, str(Path(__file__).parent.parent / "backend"))

from parser import parse_pdf, parse_docx


@pytest.fixture
def sample_pdf():
    """
    Create a sample PDF file for testing.

    Yields:
        Path to a temporary PDF file with sample content.
    """
    # Create a temporary PDF file
    with tempfile.NamedTemporaryFile(mode='wb', suffix='.pdf', delete=False) as f:
        pdf_path = f.name

    # Create a simple PDF with PyMuPDF
    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((72, 72), "This is page 1.\nIt has multiple lines.")
    page2 = doc.new_page()
    page2.insert_text((72, 72), "This is page 2.\nMore content here.")
    doc.save(pdf_path)
    doc.close()

    yield pdf_path

    # Cleanup
    if os.path.exists(pdf_path):
        os.unlink(pdf_path)


@pytest.fixture
def empty_pdf():
    """
    Create an empty PDF file for testing.

    Yields:
        Path to a temporary empty PDF file.
    """
    with tempfile.NamedTemporaryFile(mode='wb', suffix='.pdf', delete=False) as f:
        pdf_path = f.name

    # Create an empty PDF
    doc = fitz.open()
    doc.new_page()  # Empty page with no text
    doc.save(pdf_path)
    doc.close()

    yield pdf_path

    # Cleanup
    if os.path.exists(pdf_path):
        os.unlink(pdf_path)


@pytest.fixture
def sample_docx():
    """
    Create a sample DOCX file for testing.

    Yields:
        Path to a temporary DOCX file with sample content.
    """
    with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as f:
        docx_path = f.name

    # Create a simple DOCX with python-docx
    doc = Document()
    doc.add_paragraph("This is the first paragraph.")
    doc.add_paragraph("This is the second paragraph.")
    doc.add_paragraph("And here is a third one with more text.")
    doc.save(docx_path)

    yield docx_path

    # Cleanup
    if os.path.exists(docx_path):
        os.unlink(docx_path)


@pytest.fixture
def empty_docx():
    """
    Create an empty DOCX file for testing.

    Yields:
        Path to a temporary empty DOCX file.
    """
    with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as f:
        docx_path = f.name

    # Create an empty DOCX
    doc = Document()
    doc.save(docx_path)

    yield docx_path

    # Cleanup
    if os.path.exists(docx_path):
        os.unlink(docx_path)


class TestParsePDF:
    """Tests for parse_pdf function."""

    def test_parse_pdf_extracts_text(self, sample_pdf):
        """Test that parse_pdf successfully extracts text from a PDF."""
        result = parse_pdf(sample_pdf)

        assert isinstance(result, str)
        assert "This is page 1" in result
        assert "This is page 2" in result
        assert "multiple lines" in result
        assert "More content here" in result

    def test_parse_pdf_handles_empty_document(self, empty_pdf):
        """Test that parse_pdf returns empty string for empty PDF."""
        result = parse_pdf(empty_pdf)

        assert result == ""

    def test_parse_pdf_file_not_found(self):
        """Test that parse_pdf raises ValueError for non-existent file."""
        with pytest.raises(ValueError, match="File not found"):
            parse_pdf("/path/to/nonexistent/file.pdf")

    def test_parse_pdf_concatenates_pages(self, sample_pdf):
        """Test that parse_pdf concatenates multiple pages."""
        result = parse_pdf(sample_pdf)

        # Should contain content from both pages
        lines = result.split('\n')
        assert len(lines) > 0
        # Check that we have content from multiple pages
        assert any("page 1" in line for line in lines)
        assert any("page 2" in line for line in lines)


class TestParseDOCX:
    """Tests for parse_docx function."""

    def test_parse_docx_extracts_text(self, sample_docx):
        """Test that parse_docx successfully extracts text from a DOCX."""
        result = parse_docx(sample_docx)

        assert isinstance(result, str)
        assert "first paragraph" in result
        assert "second paragraph" in result
        assert "third one" in result

    def test_parse_docx_handles_empty_document(self, empty_docx):
        """Test that parse_docx returns empty string for empty DOCX."""
        result = parse_docx(empty_docx)

        assert result == ""

    def test_parse_docx_file_not_found(self):
        """Test that parse_docx raises ValueError for non-existent file."""
        with pytest.raises(ValueError, match="File not found"):
            parse_docx("/path/to/nonexistent/file.docx")

    def test_parse_docx_joins_paragraphs_with_newlines(self, sample_docx):
        """Test that parse_docx joins paragraphs with newline separators."""
        result = parse_docx(sample_docx)

        # Should have newlines between paragraphs
        lines = result.split('\n')
        assert len(lines) == 3
        assert "first paragraph" in lines[0]
        assert "second paragraph" in lines[1]
        assert "third one" in lines[2]
